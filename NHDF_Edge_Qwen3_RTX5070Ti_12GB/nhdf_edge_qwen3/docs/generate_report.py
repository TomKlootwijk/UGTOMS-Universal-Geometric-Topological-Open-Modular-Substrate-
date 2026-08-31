#!/usr/bin/env python3
"""Generate the legacy v0.1 NHDF Edge engineering report as DOCX.

This generator predates the v0.3 specification and the measured custom-codec
quality failure. See docs/README.md before regenerating or citing its output.

The DOCX is converted and visually verified by the repository build workflow.
"""
from __future__ import annotations

import csv
import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
FIG = ROOT / "figures"
MET = ROOT / "metrics"
OUT = DOCS / "NHDF_Edge_Qwen3_RTX5070Ti_Report.docx"

NAVY = "153B57"
TEAL = "0B8F8F"
PALE_TEAL = "E7F5F4"
PALE_BLUE = "EAF1F6"
PALE_AMBER = "FFF4D8"
PALE_RED = "FCE8E6"
DARK = "1F2933"
MID = "53616F"
LIGHT = "D9E2E8"
WHITE = "FFFFFF"
GRAY = "F5F7F8"


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, **edges) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if edge in edges:
            edge_data = edges[edge]
            tag = "w:" + edge
            element = tc_borders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tc_borders.append(element)
            for key in ("val", "sz", "space", "color"):
                if key in edge_data:
                    element.set(qn("w:" + key), str(edge_data[key]))


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_row_cant_split(row) -> None:
    """Keep a table row on one page when the row can fit."""
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_cell_margins(cell, top=80, start=90, bottom=80, end=90) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    run.font.name = "Inter"
    run.font.size = Pt(8)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths=None, font_size=8.3):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    set_row_cant_split(hdr)
    for idx, text in enumerate(headers):
        cell = hdr.cells[idx]
        cell.text = text
        shade(cell, NAVY)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for r in p.runs:
                r.font.name = "Inter"
                r.font.size = Pt(font_size)
                r.font.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_margins(cell)
        set_cell_border(cell, bottom={"val": "single", "sz": 5, "color": WHITE})
    for r_idx, values in enumerate(rows):
        row = table.add_row()
        set_row_cant_split(row)
        for idx, text in enumerate(values):
            cell = row.cells[idx]
            cell.text = str(text)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            shade(cell, WHITE if r_idx % 2 == 0 else GRAY)
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    run.font.name = "Noto Sans"
                    run.font.size = Pt(font_size)
                    run.font.color.rgb = RGBColor.from_string(DARK)
            set_cell_margins(cell)
            set_cell_border(cell, bottom={"val": "single", "sz": 3, "color": LIGHT})
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_callout(doc: Document, title: str, body: str, fill=PALE_BLUE, accent=NAVY) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(0.12)
    table.columns[1].width = Inches(6.7)
    shade(table.cell(0, 0), accent)
    shade(table.cell(0, 1), fill)
    cell = table.cell(0, 1)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(title)
    run.bold = True
    run.font.name = "Inter"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor.from_string(NAVY)
    p2 = cell.add_paragraph(body)
    p2.paragraph_format.space_after = Pt(0)
    for r in p2.runs:
        r.font.name = "Noto Sans"
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor.from_string(DARK)
    for c in table.rows[0].cells:
        set_cell_margins(c, top=100, start=110, bottom=100, end=110)
        set_cell_border(c, top={"val": "nil"}, bottom={"val": "nil"}, left={"val": "nil"}, right={"val": "nil"})
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_code_block(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_row_cant_split(table.rows[0])
    cell = table.cell(0, 0)
    shade(cell, "F1F4F6")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    for i, line in enumerate(text.splitlines()):
        if i:
            p.add_run().add_break()
        r = p.add_run(line)
        r.font.name = "Liberation Mono"
        r.font.size = Pt(8.3)
        r.font.color.rgb = RGBColor.from_string(DARK)
    set_cell_margins(cell, top=110, start=130, bottom=110, end=130)
    set_cell_border(cell, top={"val": "single", "sz": 4, "color": LIGHT}, bottom={"val": "single", "sz": 4, "color": LIGHT}, left={"val": "single", "sz": 4, "color": LIGHT}, right={"val": "single", "sz": 4, "color": LIGHT})
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_figure(doc: Document, filename: str, caption: str, width=6.75) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    run.add_picture(str(FIG / filename), width=Inches(width))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(8)
    for r in cap.runs:
        r.font.name = "Noto Sans"
        r.font.size = Pt(8)
        r.font.italic = True
        r.font.color.rgb = RGBColor.from_string(MID)


def add_heading(doc: Document, text: str, level=1, page_break=False):
    """Add a heading without inserting a standalone page-break paragraph.

    Using paragraph_format.page_break_before avoids occasional blank pages when
    the preceding content already ends exactly at a page boundary.
    """
    heading = doc.add_heading(text, level=level)
    heading.paragraph_format.page_break_before = bool(page_break)
    return heading


def add_body(doc: Document, text: str, bold_prefix: str | None = None):
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)
    return p


def add_bullets(doc: Document, items: list[str], level=0):
    for item in items:
        p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
        p.add_run(item)


def setup_document() -> Document:
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.65)
    sec.bottom_margin = Inches(0.62)
    sec.left_margin = Inches(0.78)
    sec.right_margin = Inches(0.78)
    sec.header_distance = Inches(0.25)
    sec.footer_distance = Inches(0.25)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Noto Sans"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Sans")
    normal.font.size = Pt(9.3)
    normal.font.color.rgb = RGBColor.from_string(DARK)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.08

    for name, size, color in (
        ("Title", 30, WHITE),
        ("Subtitle", 14, "DDEAF2"),
        ("Heading 1", 20, NAVY),
        ("Heading 2", 13, TEAL),
        ("Heading 3", 10.5, NAVY),
    ):
        style = styles[name]
        style.font.name = "Inter"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Inter")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        if name.startswith("Heading"):
            style.font.bold = True
            style.paragraph_format.keep_with_next = True
            style.paragraph_format.space_before = Pt(10 if name == "Heading 1" else 7)
            style.paragraph_format.space_after = Pt(4)
    styles["Heading 1"].paragraph_format.page_break_before = False

    for list_name in ("List Bullet", "List Bullet 2", "List Number"):
        style = styles[list_name]
        style.font.name = "Noto Sans"
        style.font.size = Pt(9.1)
        style.paragraph_format.space_after = Pt(2.5)

    header = sec.header
    p = header.paragraphs[0]
    p.text = "NHDF Edge AI | Qwen3-30B-A3B | RTX 5070 Ti Laptop 12 GB"
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for r in p.runs:
        r.font.name = "Inter"
        r.font.size = Pt(7.5)
        r.font.color.rgb = RGBColor.from_string(MID)
    footer = sec.footer
    fp = footer.paragraphs[0]
    add_page_number(fp)

    props = doc.core_properties
    props.title = "NHDF Edge AI: Qwen3-30B-A3B on a 12 GB RTX 5070 Ti Laptop GPU"
    props.subject = "Feasibility, architecture, packing format, metrics and validation plan"
    props.author = "Tom Klootwijk"
    props.keywords = "NHDF, edge AI, Qwen3, mixture of experts, quantization, RTX 5070 Ti"
    props.comments = "Conceptual research engineering report. Analytical projections are not measured benchmarks."
    return doc


def cover(doc: Document, projection: dict) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade(cell, NAVY)
    set_cell_margins(cell, top=400, start=260, bottom=400, end=260)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("NHDF Edge AI")
    r.font.name = "Inter"
    r.font.size = Pt(31)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(WHITE)
    p2 = cell.add_paragraph("Qwen3-30B-A3B on a 12 GB RTX 5070 Ti Laptop GPU")
    p2.paragraph_format.space_before = Pt(8)
    for r in p2.runs:
        r.font.name = "Inter"
        r.font.size = Pt(17)
        r.font.color.rgb = RGBColor.from_string("DDEAF2")
    p3 = cell.add_paragraph("Architecture, weight-packing format, edge execution plan, feasibility metrics and falsifiable validation program")
    p3.paragraph_format.space_before = Pt(12)
    for r in p3.runs:
        r.font.name = "Noto Sans"
        r.font.size = Pt(10.5)
        r.font.color.rgb = RGBColor.from_string(WHITE)

    doc.add_paragraph()
    add_figure(doc, "operator_pipeline.png", "The formal NHDF operator chain mapped to an implementable edge-AI weight runtime.", width=6.8)

    status = doc.add_table(rows=2, cols=3)
    status.alignment = WD_TABLE_ALIGNMENT.CENTER
    status.autofit = False
    headings = ["Feasibility", "Packed weights", "Modeled total VRAM"]
    values = ["Conditional / experimental", f"{projection['packed_weight_gb']:.2f} GB", f"{projection['projected_total_vram_gb']:.2f} GB at 8K"]
    for c, text in enumerate(headings):
        status.cell(0, c).text = text
        shade(status.cell(0, c), TEAL)
        for r in status.cell(0, c).paragraphs[0].runs:
            r.font.name = "Inter"; r.font.size = Pt(8); r.font.bold = True; r.font.color.rgb = RGBColor.from_string(WHITE)
    for c, text in enumerate(values):
        status.cell(1, c).text = text
        shade(status.cell(1, c), PALE_TEAL)
        for r in status.cell(1, c).paragraphs[0].runs:
            r.font.name = "Noto Sans"; r.font.size = Pt(9); r.font.bold = True; r.font.color.rgb = RGBColor.from_string(NAVY)
    for row in status.rows:
        for cell in row.cells:
            set_cell_margins(cell, top=110, start=100, bottom=110, end=100)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    doc.add_paragraph()
    p = doc.add_paragraph("Prepared for Tom Klootwijk")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for r in p.runs:
        r.font.name = "Inter"; r.font.size = Pt(11); r.font.bold = True; r.font.color.rgb = RGBColor.from_string(NAVY)
    p = doc.add_paragraph("Version 0.1 | 31 August 2026")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for r in p.runs:
        r.font.name = "Noto Sans"; r.font.size = Pt(9); r.font.color.rgb = RGBColor.from_string(MID)
    add_callout(
        doc,
        "Document status",
        "Conceptual research engineering report. The source-derived operator specification is implemented as a semantics-first prototype. Model quality and GPU speed have not been measured on the target laptop; all forward-looking performance figures are explicitly labeled analytical.",
        fill=PALE_AMBER,
        accent="C58B00",
    )


def main() -> None:
    analytical = json.loads((MET / "analytical_projection.json").read_text(encoding="utf-8"))
    projection = analytical["estimate"]
    smoke = json.loads((MET / "smoke_test.json").read_text(encoding="utf-8"))
    residual_rows = list(csv.DictReader((MET / "residual_fraction_sweep.csv").open(encoding="utf-8")))
    context_rows = list(csv.DictReader((MET / "context_vram_sweep.csv").open(encoding="utf-8")))

    doc = setup_document()
    cover(doc, projection)

    # Document control and contents.
    add_heading(doc, "Document control and reading guide", level=1, page_break=True)
    add_table(
        doc,
        ["Item", "Value"],
        [
            ["Design basis", "NHDF Formal Specification v0.1 by Tom Klootwijk [S1], plus the supplied prompt-chain source [S0]."],
            ["Selected model", "Qwen/Qwen3-30B-A3B-Instruct-2507 [M1]."],
            ["Target", "GeForce RTX 5070 Ti Laptop GPU, 12 GB GDDR7 capacity class [H1]."],
            ["Implemented evidence", "Packer, local zero-set projection, log-polar branch score, parity, CRC, selected-row decode, Qwen3 expert adapter, CPU tests, analytical metrics."],
            ["Not yet measured", "Complete checkpoint size, language-model quality, target CUDA kernel correctness/performance, prefill performance, power and thermal behavior."],
            ["ZIP boundary", "No upstream model weights are included. Pull and conversion happen locally."],
        ],
        widths=[1.55, 5.15],
        font_size=8.4,
    )
    add_body(doc, "Contents")
    add_table(
        doc,
        ["Section", "Purpose"],
        [
            ["1. Executive conclusion", "Decision, fit, expected use and principal risks."],
            ["2. Evidence boundary", "Separates source requirements, established facts, implementation, projections and hypotheses."],
            ["3. Model and hardware selection", "Why this model is useful and why it does not fit conventional 12 GB deployment."],
            ["4. Formal-to-runtime architecture", "Exact mapping of ELP, B0, P, RBST, K_T, Scone, Pi and U."],
            ["5. Weight format and distillation", "Low-bit codes, local zero sets, residual branches, calibration and integrity."],
            ["6. Edge execution", "Decode, prefill, Qwen3-MoE integration and memory schedule."],
            ["7. Compression and VRAM metrics", "Analytical model and scaling with residual budget/context."],
            ["8. Speed model", "Traffic roofline, sensitivity and what must be benchmarked."],
            ["9. Quality and validation", "Synthetic smoke result, ablations, fault tests and deployment gates."],
            ["10. Applications and operational guide", "Practical edge uses and commands."],
            ["Appendices", "Equations, telemetry contract, package map and references."],
        ],
        widths=[2.2, 4.5],
        font_size=8.2,
    )

    # 1 Executive conclusion.
    add_heading(doc, "1. Executive conclusion", level=1, page_break=True)
    add_callout(
        doc,
        "Feasibility verdict: conditional",
        f"The proposed mixed-precision NHDF Edge pack is analytically {projection['packed_weight_gb']:.2f} GB and the full default 8K profile is {projection['projected_total_vram_gb']:.2f} GB, leaving {projection['nominal_headroom_gb']:.2f} GB of nominal 12 GB capacity. This is enough to justify a build-and-benchmark experiment, but not enough to claim deployment before the actual laptop reports free VRAM, the complete pack is serialized, and the CUDA path is measured.",
        fill=PALE_AMBER,
        accent="C58B00",
    )
    add_body(doc, "Qwen3-30B-A3B-Instruct-2507 is the practical stress target because it combines a large 30.5B total parameter footprint with only 3.3B active parameters per token. Its 128-expert, top-8 routing makes expert-selective 2-bit compression and on-demand expert projection technically meaningful [M1][M2]. The upstream BF16 repository is 61.1 GB and the official GPTQ-Int4 checkpoint is 16.9 GB, so neither is an all-weights-resident 12 GB solution [M1][M3].")
    add_body(doc, "The implementation does not attempt to make memory literally geometric or to make a GPU execute topology directly. It translates the formal specification into a concrete data representation: each 256-weight group has a non-degenerate local zero-set constraint, residual error is log-polar encoded, a bounded one-bit correction branch is allocated by deterministic error/phase ranking, one-bit parity is a fast event, CRC32 protects each tensor file, and fused kernels reconstruct only the rows needed for a projection.")
    add_bullets(doc, [
        "Default experts: 2-bit base plus a one-bit sign residual on 15% of groups; modeled 2.321 bits/weight including metadata.",
        "Attention, embeddings and LM head: groupwise 4-bit; router and normalization weights remain FP16.",
        "Default context: 8,192 tokens with a modeled int8 KV cache of 0.403 GB.",
        "Decode traffic: 1.182 GB of packed active weights per generated token; modeled 17.0, 28.4 and 39.8 tokens/s at 3%, 5% and 7% of published bandwidth.",
        "Production requirement: fused low-bit GEMV for decode, packed GEMM or bounded layer-local dequantization for prefill, and exact CPU/GPU equivalence tests.",
    ])
    add_figure(doc, "model_size_comparison.png", "Figure 1. Upstream and projected weight footprints. The 12 GB line is a capacity ceiling, not usable free VRAM.")

    # 2 Evidence boundary.
    add_heading(doc, "2. Evidence boundary and corrections to the source narrative", level=1, page_break=True)
    add_body(doc, "The formal specification already distinguishes its normative core from speculative applications and unsupported hardware claims [S1]. This report maintains that separation. The earlier source transcript is used for vocabulary and traceability, but statements such as mathematical impossibility of crashes, instantaneous processing, storm immunity, direct Tensor Core parity, or near-zero processing cost are not adopted as engineering facts [S0].")
    add_table(
        doc,
        ["Class", "What this report treats as evidence", "Examples"],
        [
            ["Source-derived requirement", "Normative design intent from the supplied specification.", "Operator order, local zero sets, parity limits, bounded branches, causal timeline, feedback."],
            ["Established external fact", "Current official model, framework or hardware documentation.", "Qwen dimensions and file sizes; GPU memory/bandwidth/TGP; official expert tensor layout."],
            ["Implemented behavior", "Code executed in this package's CPU environment.", "Bit packing, reconstruction, local field residual, parity, CRC, selected rows, expert adapter, 41 tests."],
            ["Analytical projection", "Formula driven by declared assumptions; not a measurement.", "9.23 GB weights, 11.28 GB VRAM, bandwidth-based token-rate sensitivity."],
            ["Unverified hypothesis", "Requires full model and target hardware experiments.", "Quality retention, target kernel efficiency, useful speed, thermal stability, application advantage."],
        ],
        widths=[1.25, 2.65, 2.85],
        font_size=7.9,
    )
    add_heading(doc, "2.1 Required non-degenerate reading of SDF zero", level=2)
    add_body(doc, "A single global field F(x) = 0 for every x has zero gradient wherever differentiable. It cannot supply a normal, distance, inside/outside distinction or restoration direction. The formal specification therefore requires local constraints F_i(z_i, t) = 0, one per valid cell [S1]. In this weight format, cell i is a quantization group and its local field is the activation-weighted mean reconstruction residual.")
    add_code_block(doc, "F_i = sum_j h_ij * (w_ij - w_hat_ij) / sum_j h_ij\nChoose mu_i so that F_i = 0 for every packed group i.")
    add_heading(doc, "2.2 Parity is an event, not error correction", level=2)
    add_body(doc, "XOR parity detects an odd number of bit inversions and misses every even number. It neither locates nor repairs damage. NHDF Edge stores the one-bit gate for source conformance and fast diagnostics, while CRC32 protects the complete safetensors file. A deployable corruption-recovery claim would require ECC, redundancy or recomputation [S1].")
    add_callout(doc, "Safety boundary", "The runtime must fail closed on CRC mismatch, missing tensors, impossible geometry, insufficient free VRAM or unresolved meta parameters. Topology is not hardware resilience.", fill=PALE_RED, accent="B42318")

    # 3 Model/hardware.
    add_heading(doc, "3. Model and target selection", level=1, page_break=True)
    add_heading(doc, "3.1 Why Qwen3-30B-A3B-Instruct-2507", level=2)
    add_table(
        doc,
        ["Property", "Value", "Architectural consequence"],
        [
            ["Total / active parameters", "30.5B / 3.3B", "Large stored state, much smaller conditional compute per token."],
            ["Layers", "48", "All-layer weight residency matters; layer offload would penalize every token."],
            ["Experts", "128 total, 8 active/token", "2-bit expert compression yields most of the size reduction; selected expert rows can be projected."],
            ["Expert intermediate", "768", "Each expert has fused gate/up and down matrices with predictable row intervals."],
            ["Hidden size", "2,048", "Group size 256 divides the primary input width exactly."],
            ["Attention", "32 query heads, 4 KV heads, head dimension 128", "Grouped-query attention keeps the KV cache relatively small."],
            ["Context", "262,144 native maximum", "The laptop profile intentionally caps default deployment at 8K; native maximum is not a 12 GB target."],
        ],
        widths=[1.35, 1.4, 4.0],
        font_size=8.1,
    )
    add_body(doc, "Hugging Face's official Qwen3-MoE implementation stores each expert collection as two 3-D parameters: gate_up_proj[num_experts, 2*intermediate, hidden] and down_proj[num_experts, hidden, intermediate]. Its reference forward groups tokens by selected expert, performs the two projections and accumulates weighted outputs [M2]. The included adapter preserves this contract rather than flattening experts into unrelated linear modules.")
    add_heading(doc, "3.2 Target GPU facts and practical constraints", level=2)
    add_table(
        doc,
        ["Target fact", "Published value", "Use in this analysis"],
        [
            ["VRAM", "12 GB GDDR7", "Hard nominal capacity; actual free memory must be measured after display/driver allocation."],
            ["Memory interface / bandwidth", "192-bit / 672 GB/s", "Roofline input for decode traffic sensitivity, never treated as sustained application bandwidth."],
            ["CUDA cores", "5,888", "General compute context; not used to infer token rate."],
            ["AI throughput", "992 AI TOPS", "Not used for the low-bit GEMV estimate because format/kernel utilization is unknown."],
            ["Power range", "60-115 W", "Laptop configuration and thermals can materially change sustained performance."],
            ["Architecture", "Blackwell; fifth-generation Tensor Cores", "Use only through supported CUDA/PyTorch interfaces. No direct topology or parity execution claim."],
        ],
        widths=[1.45, 1.55, 3.75],
        font_size=8.1,
    )
    add_body(doc, "The implementation deliberately detects the runtime compute capability instead of baking a speculative PTX sequence into the format. CUDA 12.8 introduced full Blackwell toolkit support, but the exact compiler, PyTorch build, driver and architecture spelling on the laptop must be checked together [H2].")

    # 4 formal runtime architecture.
    add_heading(doc, "4. Formal-to-runtime architecture", level=1, page_break=False)
    add_figure(doc, "operator_pipeline.png", "Figure 2. Normative order retained in the weight-runtime mapping.")
    add_table(
        doc,
        ["Operator", "Formal role [S1]", "NHDF Edge realization", "Testable invariant"],
        [
            ["ELP", "Jitter to log-polar address", "Residual L2 magnitude to rho; deterministic sine/cosine projections to theta; one-byte bin for selected groups.", "Address is deterministic and bounded to declared bins."],
            ["B0", "Project a candidate to a local zero-level set", "Solve one FP16 group mean so the weighted residual average is zero.", "max |F_i| <= declared tolerance."],
            ["P", "One-bit data/topology event", "XOR parity of base and residual payload bytes; orientation is a separate diagnostic; file CRC is authoritative.", "Stored parity exactly recomputes; even-fault blind spot documented."],
            ["RBST", "Bounded branch-node bifurcation", "Stable sort of group error times phase-curvature factor; allocate exactly the highest-scoring residual budget.", "Selected count <= finite budget; deterministic ties."],
            ["K_T", "Forward causal propagation", "Monotonic tensor conversion order, manifest generation and autoregressive token/cache sequence.", "No future-generation read; bounded buffers carry generation tags where used."],
            ["Scone", "Analytic sweep / extrusion", "Decode low-bit codes plus local and residual terms only for requested weight rows or tiles.", "No full dense weight is required for batch-one projection."],
            ["Pi", "Observable projection", "Embedding gather, GEMV/GEMM, expert projection and logits.", "CPU/GPU outputs agree within declared numerical tolerance."],
            ["U", "Self-referential state update", "Manifest/telemetry, residual metrics, optional Hessian calibration and next conversion/runtime state.", "Run emits traceable configuration and failure state."],
        ],
        widths=[0.55, 1.35, 3.25, 1.55],
        font_size=7.25,
    )
    add_heading(doc, "4.1 What is self-referential here", level=2)
    add_body(doc, "The current packed state is both operand and update target: the residual produced by reconstructing a source tensor determines log-polar address, parity, branch allocation and the serialized next state. At runtime, outputs and residual telemetry can select recalibration or a new residual budget. This is a bounded closed state transition, not unrestricted recursive memory growth.")
    add_code_block(doc, "L_(n+1) = U(L_n, Pi(Scone(K_T(RBST(P(B0(ELP(x_n; L_n))))))))")

    # 5 format/distillation.
    add_heading(doc, "5. Weight representation and reconstruction distillation", level=1, page_break=False)
    add_heading(doc, "5.1 Tensor policies", level=2)
    add_table(
        doc,
        ["Tensor class", "Default", "Reason"],
        [
            ["MoE expert gate/up and down", "2-bit base + 15% one-bit residual", "Experts dominate stored parameters and published MoE research supports investigating ultra-low-bit expert-only compression [Q3]."],
            ["Attention matrices", "4-bit groupwise", "Attention is active on every token and more sensitive than dormant experts; 4-bit reduces risk."],
            ["Token embedding and LM head", "4-bit groupwise", "Large matrices, but vocabulary/logit quality is sensitive; kept above 2-bit."],
            ["Router", "FP16", "Small and directly controls expert selection."],
            ["RMSNorm and other 1-D state", "FP16", "Negligible capacity gain from lower precision; preserves scaling stability."],
            ["KV cache", "Modeled int8", "Separate runtime choice; must be validated against a concrete cache implementation."],
        ],
        widths=[2.0, 1.8, 2.95],
        font_size=8.0,
    )
    add_heading(doc, "5.2 Group reconstruction", level=2)
    add_body(doc, "For group i and element j, the reconstructed weight is a local mean plus a low-bit level and an optional sign residual. The mean is solved after the residual branch is chosen, so the local field remains at zero within floating-point tolerance.")
    add_code_block(doc, "w_hat_ij = mu_i + s_i * level(q_ij) + m_i * r_i * sign_ij\nq_ij: 2-bit or 4-bit code\nm_i: residual-mask bit for group i\nmu_i, s_i, r_i: FP16 local parameters")
    add_body(doc, "The converter uses alternating least squares for the base scale and mean. With an optional per-input-channel second moment h_j, both the local field and branch error become activation weighted. This is post-training reconstruction distillation: the source weight/teacher activations define the loss, but the package does not claim full logit-level knowledge distillation.")
    add_heading(doc, "5.3 Log-polar address and phase-accelerated branch score", level=2)
    add_code_block(doc, "rho_i = log(1 + gamma * ||residual_i||_2)\ntheta_i = atan2(<residual_i, sine_basis>, <residual_i, cosine_basis>)\nscore_i = weighted_MSE_i * (1 + k_phi * |Delta2(theta)_i|)")
    add_body(doc, "A stable descending sort selects the exact finite branch budget. The golden-ratio concept remains available as a scheduling or geometric scale in the general NHDF specification, but it is not inserted as an unexplained constant into this weight codec. Equal-budget ablations must determine whether any phi-based schedule adds value.")
    add_heading(doc, "5.4 Storage layout and integrity", level=2)
    add_table(
        doc,
        ["Component", "Storage", "Expert contribution"],
        [
            ["Base codes", "2 bits/weight", "2.000 bpp"],
            ["Residual signs", "1 bit/weight on selected fraction f", "0.150 bpp at f=0.15"],
            ["Local mean + scale", "2 FP16 values per 256 weights", "0.125 bpp"],
            ["Residual scale", "1 FP16 per selected group", "0.0094 bpp"],
            ["Log-polar address", "1 byte per selected group", "0.0047 bpp"],
            ["Mask, rank prefix, parity", "three 32-bit words per 32 groups", "0.0117 bpp"],
            ["Header/alignment allowance", "analytical allowance", "0.020 bpp"],
            ["Total", "modeled", f"{projection['expert_effective_bpp']:.4f} bpp"],
        ],
        widths=[2.1, 2.8, 1.6],
        font_size=8.1,
    )
    add_body(doc, "Each packed tensor is one safetensors file. The manifest records original shape, group geometry, policy, reconstruction statistics, per-file size and CRC32. Per-group parity supports local diagnostics; CRC32 catches the even-count faults that parity can miss. The complete converted model is not included in the ZIP.")

    # 6 runtime.
    add_heading(doc, "6. Edge execution architecture", level=1, page_break=False)
    add_heading(doc, "6.1 Decode path", level=2)
    add_body(doc, "Autoregressive batch-one decode is the first optimization target because each token performs matrix-vector products and rereads active weights. The supplied CUDA extension decodes 2-bit/4-bit codes, loads local mean/scale, uses a 32-group mask plus prefix-popcount to locate residual signs/scales, and accumulates directly into the output row. It supports contiguous row intervals so one selected expert can be projected without reconstructing all 128 experts.")
    add_code_block(doc, "for output row r in requested interval:\n    for input column c in parallel:\n        group = r * groups_per_row + c / 256\n        w = mean[group] + scale[group] * level(code[r,c])\n        if residual_mask[group]:\n            rank = prefix[word] + popcount(bits before group)\n            w += residual_scale[rank] * residual_sign[rank,c]\n        accumulator += x[c] * w")
    add_callout(doc, "Kernel status", "The CUDA source is included but could not be compiled in this CPU-only artifact environment. It is a clear experimental GEMV, not a tuned Blackwell kernel. Target compilation, Compute Sanitizer, numerical equivalence and Nsight profiling are mandatory.", fill=PALE_AMBER, accent="C58B00")
    add_heading(doc, "6.2 Prefill path", level=2)
    add_body(doc, "Prompt prefill is matrix-matrix work. Reusing a batch-one GEMV for every prompt token would reread weights and waste available matrix hardware. Version 0.1 therefore treats prefill optimization as an explicit next gate. Two bounded alternatives are proposed: a fused packed GEMM, or layer-local dequantization into the 0.75 GB workspace followed by cuBLAS GEMM. Hidden allocations are prohibited because they invalidate the fit model.")
    add_heading(doc, "6.3 Full Qwen3 integration", level=2)
    add_bullets(doc, [
        "Build the Hugging Face causal LM on the meta device, so no dense BF16 allocation occurs.",
        "Replace every ordinary Linear and Embedding with NHDF packed modules.",
        "Replace each Qwen3MoeExperts object with a paired packed gate_up/down adapter that preserves official router inputs and weighted accumulation.",
        "Load router and normalization parameters as FP16.",
        "Reject the model if any meta parameter remains or a pack entry is unconsumed.",
        "Keep Transformers-specific integration isolated so version drift does not change the serialized format.",
    ])
    add_heading(doc, "6.4 Memory schedule", level=2)
    add_figure(doc, "vram_budget.png", "Figure 3. Default analytical VRAM budget. Actual free VRAM, allocator fragmentation and temporary buffers must be measured.")

    # 7 metrics.
    add_heading(doc, "7. Compression and VRAM metrics", level=1, page_break=False)
    add_callout(doc, "Metric status", "Every number in Sections 7 and 8 is formula-derived from official model/hardware data and declared format assumptions. The only executed quality result is the small synthetic smoke experiment in Section 9.", fill=PALE_BLUE, accent=TEAL)
    add_heading(doc, "7.1 Default result", level=2)
    add_table(
        doc,
        ["Metric", "Value", "Interpretation"],
        [
            ["Expert effective precision", f"{projection['expert_effective_bpp']:.4f} bpp", "Includes base, residual and modeled metadata."],
            ["Sensitive matrix precision", f"{projection['sensitive_effective_bpp']:.4f} bpp", "4-bit code plus mean/scale/parity and allowance."],
            ["Packed weights", f"{projection['packed_weight_gb']:.3f} GB / {projection['packed_weight_gib']:.3f} GiB", "Expected serialized/in-memory payload before measured file headers and allocator effects."],
            ["Compression vs BF16 repository", f"{projection['compression_vs_bf16']:.2f}x", "61.1 GB divided by projected packed bytes."],
            ["Compression vs official GPTQ-Int4", f"{projection['compression_vs_official_int4']:.2f}x", "16.9 GB divided by projected packed bytes."],
            ["8K int8 KV cache", f"{projection['kv_cache_gb']:.3f} GB", "Batch one, K and V, 48 layers, 4 KV heads, head dimension 128."],
            ["Total default profile", f"{projection['projected_total_vram_gb']:.3f} GB", "Weights + KV + 0.75 GB workspace + 0.90 GB reserve."],
            ["Nominal headroom", f"{projection['nominal_headroom_gb']:.3f} GB", "Fragile; actual free VRAM may be lower."],
        ],
        widths=[2.05, 1.65, 3.0],
        font_size=8.0,
    )
    add_heading(doc, "7.2 Residual-budget scaling", level=2)
    table_rows = []
    for row in residual_rows:
        table_rows.append([
            f"{100*float(row['residual_fraction']):.0f}%",
            f"{float(row['expert_bpp']):.3f}",
            f"{float(row['packed_weight_gb']):.3f}",
            f"{float(row['total_vram_gb']):.3f}",
            "yes" if row['fits_12gb'].lower() == 'true' else "no",
        ])
    add_table(doc, ["Residual groups", "Expert bpp", "Weight GB", "Total VRAM GB", "Nominal fit"], table_rows, widths=[1.25, 1.25, 1.35, 1.45, 1.25], font_size=8.0)
    add_figure(doc, "residual_tradeoff.png", "Figure 4. Residual budget increases size while improving a synthetic reconstruction proxy. The synthetic curve is not model accuracy.")
    add_heading(doc, "7.3 Context scaling", level=2)
    context_table = []
    for row in context_rows:
        context_table.append([
            f"{int(row['context_tokens']):,}",
            f"{float(row['kv_cache_gb']):.3f}",
            f"{float(row['total_vram_gb']):.3f}",
            f"{float(row['headroom_gb']):.3f}",
            "yes" if row['fits_12gb'].lower() == 'true' else "no",
        ])
    add_table(doc, ["Context", "KV GB", "Total GB", "Headroom GB", "Nominal fit"], context_table, widths=[1.2, 1.2, 1.35, 1.45, 1.25], font_size=8.0)
    add_figure(doc, "context_vram.png", "Figure 5. Int8 KV-cache pressure. The 8K default is a deployment choice, not the model's advertised maximum.")

    # 8 speed.
    add_heading(doc, "8. Execution-speed scaling model", level=1, page_break=False)
    add_heading(doc, "8.1 Active weight traffic", level=2)
    add_body(doc, "A routed token uses 8/128 of the expert parameters in every MoE layer, while attention, routers and the LM head remain active. Applying the packed bits-per-parameter to those active categories gives 1.182 GB of weight traffic per generated token before cache effects, duplicate reads, activation traffic and kernel metadata. The theoretical memory roofline is therefore 672 / 1.182 = 568 tokens/s. No general-purpose implementation reaches this roofline.")
    add_code_block(doc, "active_expert_params = 28.991B * 8 / 128\nactive_bytes = active_expert_params * expert_bpp/8\n             + (attention + lm_head) * sensitive_bpp/8\n             + (router + norms) * 2 bytes\nroofline_tps = 672 GB/s / active_bytes")
    add_heading(doc, "8.2 Sensitivity, not forecast", level=2)
    add_table(
        doc,
        ["Effective share of published bandwidth", "Modeled decode rate", "Meaning"],
        [
            ["2%", f"{projection['decode_tps_by_efficiency']['2%']:.1f} tok/s", "Low-efficiency reference or heavy overhead."],
            ["3%", f"{projection['decode_tps_by_efficiency']['3%']:.1f} tok/s", "Lower engineering sensitivity point."],
            ["5%", f"{projection['decode_tps_by_efficiency']['5%']:.1f} tok/s", "Mid sensitivity point, not an expected benchmark."],
            ["7%", f"{projection['decode_tps_by_efficiency']['7%']:.1f} tok/s", "Aggressive but useful optimization target."],
            ["10%", f"{projection['decode_tps_by_efficiency']['10%']:.1f} tok/s", "Very strong effective packed-weight throughput."],
        ],
        widths=[2.25, 1.4, 3.0],
        font_size=8.2,
    )
    add_figure(doc, "decode_sensitivity.png", "Figure 6. Linear bandwidth sensitivity. Actual kernels may be compute-, launch-, cache-, routing- or power-limited.")
    add_heading(doc, "8.3 Measurements required", level=2)
    add_bullets(doc, [
        "Kernel median and p10/p90 latency for 2-bit, 4-bit and residual/no-residual shapes.",
        "Effective bytes read per projection, global-memory transactions, L2 hit rate and achieved bandwidth.",
        "Time to first token, prefill tokens/s and steady-state decode tokens/s separately.",
        "Peak allocated and reserved VRAM, including model load and prompt growth.",
        "Power, clock, temperature and thermal-throttling behavior across the laptop's 60-115 W configurations.",
        "Five or more post-warm-up runs with fixed prompts and exact software versions.",
    ])

    # 9 quality validation.
    add_heading(doc, "9. Reconstruction quality, model quality and falsification", level=1, page_break=False)
    add_heading(doc, "9.1 Executed synthetic smoke result", level=2)
    add_body(doc, f"The package executed a deterministic 512 x 1,024 synthetic matrix experiment with low-rank structure, noise and sparse outliers. Relative to plain groupwise 2-bit, the default 15% residual branch reduced reconstruction MSE by {smoke['default_mse_reduction_percent']:.1f}% while maintaining a maximum local zero-set residual of {smoke['default_zero_set_max_abs']:.2e}; stored parity recomputed successfully. The realized in-memory test payload was {smoke['default_effective_bpp']:.3f} bpp. Four-bit reconstruction had lower MSE, as expected.")
    add_callout(doc, "Do not over-interpret this result", "A synthetic weight matrix does not measure language-model perplexity, task accuracy, routing stability or downstream behavior. It verifies that the codec's correction branch changes the intended error metric and that the local constraint/parity machinery works.", fill=PALE_AMBER, accent="C58B00")
    add_heading(doc, "9.2 Minimum model-level comparison", level=2)
    add_table(
        doc,
        ["Baseline / ablation", "Question"],
        [
            ["BF16 source", "Upper reference for logits, routes and quality."],
            ["Official GPTQ-Int4", "Existing larger-than-12-GB quality/format reference."],
            ["Plain 2-bit groupwise", "Does the NHDF residual branch help at its extra memory cost?"],
            ["Plain 3-bit or equal-size codebook", "Can a simpler representation match the same bytes?"],
            ["NHDF without B0", "Does local zero-set mean projection add value beyond an ordinary zero point?"],
            ["NHDF without log-polar / Delta2 phase", "Does the source-derived ranking outperform pure weighted MSE?"],
            ["Residual fractions 0-30%", "Where is the quality/size knee on the exact Qwen checkpoint?"],
            ["Parity disabled", "Does the event gate matter to runtime recovery/telemetry, apart from CRC?"],
        ],
        widths=[2.45, 4.2],
        font_size=8.1,
    )
    add_heading(doc, "9.3 Required quality metrics", level=2)
    add_bullets(doc, [
        "Held-out perplexity or another declared language-model likelihood metric.",
        "Instruction, reasoning and code tasks matching intended deployment.",
        "Long-context retrieval at 4K and 8K.",
        "Logit cosine similarity, top-k token agreement and greedy output drift.",
        "Expert-route divergence: fraction of tokens whose top-8 expert set/order differs from BF16.",
        "Per-layer activation error to locate tensors that need a higher residual budget or precision.",
    ])
    add_heading(doc, "9.4 Fault and resource tests", level=2)
    add_bullets(doc, [
        "Flip one payload bit: parity and CRC should detect it.",
        "Flip two bits in one parity group: demonstrate parity's blind spot while CRC fails.",
        "Corrupt/delete a tensor, alter geometry or leave a meta parameter: loader must fail closed.",
        "Cross residual-mask boundaries at groups 31/32 and expert row boundaries.",
        "Reduce free VRAM below the projected requirement: doctor/loader must stop before partial activation.",
        "Run long prompts and repeated generations to expose hidden allocation growth or driver reset risk.",
    ])

    # 10 applications/guide.
    add_heading(doc, "10. Edge applications and operational guide", level=1, page_break=False)
    add_heading(doc, "10.1 Practical edge applications", level=2)
    add_table(
        doc,
        ["Application", "Why the selected model/architecture is relevant", "Primary validation metric"],
        [
            ["Private local coding assistant", "Large sparse instruction model without cloud data transfer; batch-one decode is central.", "Code task pass rate, TTFT, decode rate, VRAM."],
            ["Offline document and RAG analyst", "Local reasoning over sensitive corpora; 8K context is a useful initial bound.", "Retrieval answer quality, groundedness, 8K stability."],
            ["Field-service copilot", "Low-connectivity laptop operation; model remains resident between tasks.", "Power-limited sustained speed, reliability, domain accuracy."],
            ["Research/notebook agent", "More capability than smaller dense models may provide while keeping all experts locally stored.", "Tool-call accuracy, latency, reproducibility."],
            ["Local multi-agent orchestration", "One resident model can serve specialized prompts without network dependency.", "Concurrent memory, scheduling latency, safety constraints."],
            ["Smaller-model efficiency profile", "Same codec can trade residual budget for longer context or lower power even when fit is not the problem.", "Quality per byte/joule versus standard 4-bit."],
        ],
        widths=[1.65, 3.25, 1.85],
        font_size=7.75,
    )
    add_heading(doc, "10.2 Local workflow", level=2)
    add_code_block(doc, "# 1. Install and test\npython -m pip install -e \".[dev]\"\npytest -q\nnhdf-edge smoke\n\n# 2. Pull source weights (not included in ZIP)\npython scripts/pull_model.py --output models/Qwen3-30B-A3B-Instruct-2507\n\n# 3. Stream-convert\nnhdf-edge pack models/Qwen3-30B-A3B-Instruct-2507 packs/qwen3-nhdf \\\n  --config configs/qwen3_30b_a3b_edge12.yaml\n\n# 4. Verify every file and group parity\nnhdf-edge verify packs/qwen3-nhdf --parity-all\n\n# 5. Build target CUDA extension and inspect free VRAM\nexport TORCH_CUDA_ARCH_LIST=\"12.0\"\npython setup_cuda.py build_ext --inplace\nnhdf-edge doctor --config configs/qwen3_30b_a3b_edge12.yaml")
    add_heading(doc, "10.3 Benchmark workflow", level=2)
    add_code_block(doc, "python scripts/benchmark_kernel.py --bits 2 --residual-fraction 0.15\npython scripts/benchmark_kernel.py --bits 4 --residual-fraction 0\npython scripts/benchmark_model.py packs/qwen3-nhdf --max-new-tokens 64 \\\n  --output metrics/local/model_benchmark.json")
    add_body(doc, "The full BF16 teacher does not fit the target GPU. Optional activation calibration therefore needs CPU offload and enough system RAM, or a larger temporary machine. The data-free converter remains usable and should be tested first.")

    # 11 engineering gates and risk register.
    add_heading(doc, "11. Engineering gates, risks and decision criteria", level=1, page_break=False)
    add_table(
        doc,
        ["Gate", "Exit evidence", "Current status"],
        [
            ["G0 semantic codec", "Deterministic CPU tests, local residual tolerance, CRC/parity failure tests.", "Included; 41 tests passed in artifact environment."],
            ["G1 CUDA equivalence", "Build, sanitizer, row-offset tests, numerical error and kernel benchmark JSON.", "Source included; target run required."],
            ["G2 full conversion", "Complete manifest, actual bytes, no missing tensors, policy audit.", "Requires local source checkpoint."],
            ["G3 model correctness", "No meta state, reference logits/routes, deterministic greedy prompts.", "Loader included; full pack and CUDA required."],
            ["G4 quality/ablations", "Predeclared model metrics against equal-budget baselines.", "Not run."],
            ["G5 laptop performance", "TTFT, prefill/decode, VRAM, bandwidth, power and thermals across profiles.", "Not run."],
            ["G6 prefill optimization", "Packed GEMM or bounded dequantize-GEMM within workspace.", "Design target."],
        ],
        widths=[1.2, 4.0, 1.55],
        font_size=7.9,
    )
    add_heading(doc, "11.1 Risk register", level=2)
    add_table(
        doc,
        ["Risk", "Impact", "Mitigation / falsifier"],
        [
            ["Actual pack exceeds projection", "No safe 12 GB fit.", "Measure serialized categories; lower residual budget, context or choose a smaller model."],
            ["2-bit expert quality loss", "Model unusable despite fit.", "Activation-weighted calibration, per-layer residual budgets, 3-bit/equal-size baseline; reject if thresholds fail."],
            ["Router sensitivity", "Different experts and cascading logit error.", "Keep router FP16; measure route divergence; selectively raise nearby tensors."],
            ["Kernel efficiency below useful threshold", "Slow generation.", "Nsight profile, vectorized loads, expert compaction, optimized packed GEMV/GEMM; compare AQLM/other kernels."],
            ["Prefill bottleneck", "Poor TTFT on realistic documents.", "Layer-local dequantize-GEMM or packed GEMM; measure separately from decode."],
            ["VRAM reserve too optimistic", "OOM/driver reset.", "Use doctor, internal-display tests, configurable context/workspace, fail-before-load."],
            ["Thermal throttling", "Unstable sustained rate.", "Test 60/balanced/115 W profiles with temperature and clock telemetry."],
            ["Framework drift", "Loader breaks after update.", "Pin tested versions, keep adapter isolated, maintain fixed model vectors."],
            ["Parity misunderstood as correction", "False integrity claims.", "CRC/ECC boundary in code/report; explicit two-bit fault test."],
        ],
        widths=[2.0, 1.75, 3.0],
        font_size=7.6,
    )
    add_heading(doc, "11.2 Go/no-go criteria", level=2)
    add_bullets(doc, [
        "GO to optimization only if the complete pack loads with at least the predeclared safe free-VRAM margin and passes CRC/parity/model-structure checks.",
        "GO to application trials only if model quality meets a declared threshold and NHDF beats a simpler equal-size baseline on at least one relevant metric.",
        "NO-GO for the 30B target if prefill/decode remains below the application's minimum useful latency after reasonable kernel work; reuse the codec on a smaller model instead.",
        "NO-GO for any resilience claim unless controlled fault injection demonstrates the specific detection/recovery mechanism.",
    ])

    # 12 conclusion.
    add_heading(doc, "12. Conclusion", level=1, page_break=False)
    add_body(doc, "The formal NHDF specification can be translated into a coherent edge-AI codec without relying on physically impossible or unsupported claims. Its strongest implementable idea is the ordered closure of local validity, compact addressing, a bounded event-driven correction branch, causal execution and feedback. In this package, that becomes a mixed-precision MoE weight representation with explicit invariants and a targetable fused decode kernel.")
    add_body(doc, f"The selected Qwen3 model is a credible challenge: a conventional 61.1 GB BF16 or 16.9 GB official GPTQ-Int4 checkpoint does not fit 12 GB, whereas the default pack is projected at {projection['packed_weight_gb']:.2f} GB and the complete 8K profile at {projection['projected_total_vram_gb']:.2f} GB. The fit is narrow but sufficient to justify a controlled experiment. The projected {projection['decode_tps_by_efficiency']['3%']:.1f}-{projection['decode_tps_by_efficiency']['7%']:.1f} token/s sensitivity band is not evidence of achieved speed; it is a benchmark target tied to an explicit traffic model.")
    add_callout(doc, "Recommended next action", "Run a partial conversion and target CUDA microbenchmark first. Only after CPU/GPU equivalence and measured effective bandwidth should the full 61.1 GB source checkpoint be converted. Then evaluate actual pack bytes, model quality and 8K memory behavior before investing in prefill optimization.", fill=PALE_TEAL, accent=TEAL)

    # Appendices.
    add_heading(doc, "Appendix A. Analytical equations", level=1, page_break=False)
    add_heading(doc, "A.1 Effective expert bits per parameter", level=2)
    add_code_block(doc, "b_expert = b_base + f\n         + 32/g                 # FP16 mean + FP16 scale\n         + 16f/g + 8f/g         # residual scale + log-polar address\n         + 3/g                  # mask + prefix + parity\n         + 0.02                 # header/alignment allowance\n\nDefault: b_base=2, f=0.15, g=256 -> 2.32078125 bpp")
    add_heading(doc, "A.2 Packed weights", level=2)
    add_code_block(doc, "M_weights = N_expert * b_expert/8\n          + N_sensitive * b_sensitive/8\n          + N_raw * 2 bytes\n\nN_expert = 28,991,029,248\nN_sensitive = total - expert - router - norms\nN_raw = router + norms")
    add_heading(doc, "A.3 KV cache", level=2)
    add_code_block(doc, "M_KV = batch * context * 2(K,V) * layers * kv_heads * head_dim * bytes_per_element\nAt batch=1, context=8192, layers=48, kv_heads=4, head_dim=128, int8 -> 402,653,184 bytes")
    add_heading(doc, "A.4 Decode traffic and roofline", level=2)
    add_code_block(doc, "active_expert = N_expert * 8/128\nB_token = active_expert*b_expert/8 + active_sensitive*b_sensitive/8 + active_raw*2\nroofline = published_bandwidth / B_token\nmodeled_tps(efficiency e) = e * roofline")

    add_heading(doc, "Appendix B. Telemetry and manifest contract", level=1, page_break=False)
    add_body(doc, "Every model-level run should record enough information to reproduce both numerical and systems behavior. The following is a minimum contract; application adapters may add focus, task or safety residuals.")
    add_code_block(doc, "{\n  model_repo, model_revision, pack_format, config_hash,\n  tensor_count, actual_pack_bytes, crc_failures, parity_failures,\n  torch_version, transformers_version, cuda_runtime, driver, device,\n  free_vram_before_load, peak_allocated, peak_reserved,\n  context, batch, prompt_hash, seed, power_mode,\n  ttft, prefill_tps, decode_tps, temperature, power, clock,\n  quality_metrics, route_divergence, saturation, overflow, status\n}")
    add_table(
        doc,
        ["Manifest field", "Purpose"],
        [
            ["format / created_utc / source_model", "Version and origin."],
            ["config", "Complete model, target and packing assumptions."],
            ["tensor file / CRC32 / bytes", "Integrity and actual storage."],
            ["original shape / rows / columns / groups", "Decode geometry and bounds."],
            ["policy", "Raw/quantized, bit width, group size, residual fraction and phase parameters."],
            ["stats", "Original parameters, effective bpp, residual fraction, rho range, local-field residual and reconstruction error."],
        ],
        widths=[2.25, 4.45],
        font_size=8.1,
    )

    add_heading(doc, "Appendix C. Delivered package map", level=1, page_break=True)
    add_table(
        doc,
        ["Path", "Contents"],
        [
            ["README.md", "Feasibility summary, commands, evidence boundary and references."],
            ["configs/", "Default Qwen3/12-GB packing profile."],
            ["src/nhdf_edge/", "Operators, quantizer, pack format, converter, metrics, CLI, calibration and runtime."],
            ["csrc/ + setup_cuda.py", "Experimental fused packed GEMV and selected-row decode extension."],
            ["scripts/", "Model pull, calibration, generation, kernel/model benchmarks and report metrics."],
            ["tests/", "41 CPU tests covering codec, storage, metrics and runtime semantics."],
            ["metrics/", "Analytical JSON/CSV and deterministic synthetic smoke results."],
            ["figures/", "Report charts generated from the saved metrics."],
            ["docs/", "This report, implementation gates and validation protocol."],
            ["sources/", "NHDF Formal Specification v0.1 and source checksums; no model weights."],
        ],
        widths=[2.0, 4.7],
        font_size=8.2,
    )

    add_heading(doc, "References", level=1, page_break=True)
    refs = [
        ("[S0]", "Supplied 30-page prompt-chain transcript: original vocabulary, applications and GPU discussion. Reviewed for traceability; not redistributed in the ZIP because it contains personal identifiers and superseded claims."),
        ("[S1]", "Tom Klootwijk. Non-Euclidean Holographic Data Fields: A Formal Operator Specification for a Parity-Conditioned Kinematic Foliation and Implicit Holographic Co-Processor. Version 0.1, 31 August 2026. Included in sources/."),
        ("[M1]", "Qwen Team. Qwen/Qwen3-30B-A3B-Instruct-2507 model card, configuration and checkpoint tree. Hugging Face, accessed 31 August 2026."),
        ("[M2]", "Hugging Face Transformers. modeling_qwen3_moe.py: official Qwen3MoeExperts tensor layout and forward contract, accessed 31 August 2026."),
        ("[M3]", "Qwen Team. Qwen3-30B-A3B-GPTQ-Int4 checkpoint tree, 16.9 GB, Hugging Face, accessed 31 August 2026."),
        ("[H1]", "NVIDIA. GeForce RTX 50 Series Laptop GPU specifications: RTX 5070 Ti Laptop GPU, 12 GB GDDR7, 672 GB/s, 5,888 CUDA cores, 992 AI TOPS, 60-115 W. Accessed 31 August 2026."),
        ("[H2]", "NVIDIA. CUDA Toolkit 12.8 release notes and Blackwell support documentation."),
        ("[Q1]", "V. Egiazarian et al. Extreme Compression of Large Language Models via Additive Quantization. arXiv:2401.06118, 2024."),
        ("[Q2]", "A. Tseng et al. QuIP#: Even Better LLM Quantization with Hadamard Incoherence and Lattice Codebooks. arXiv:2402.04396, 2024."),
        ("[Q3]", "Y. J. Kim, R. Fahim, H. H. Awadalla. Mixture of Quantized Experts. arXiv:2310.02410, 2023."),
        ("[Q4]", "Fast Inference of Mixture-of-Experts Language Models with Offloading. arXiv:2312.17238, 2023. Notes that practical 2-bit formats carry metadata and keep embeddings, logits, gates and norms at higher precision."),
    ]
    add_table(doc, ["ID", "Source"], [[a, b] for a, b in refs], widths=[0.65, 6.0], font_size=7.8)
    add_callout(doc, "Citation note", "External sources establish model dimensions, checkpoint sizes, hardware specifications and low-bit precedent. None validates the NHDF format, exact Qwen quality or target laptop speed; those remain experimental obligations.", fill=PALE_BLUE, accent=NAVY)

    # Footer on last page marker.
    p = doc.add_paragraph("End of report")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs:
        r.font.name = "Inter"; r.font.size = Pt(8); r.font.italic = True; r.font.color.rgb = RGBColor.from_string(MID)

    DOCS.mkdir(exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
